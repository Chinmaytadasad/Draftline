import os
from docx import Document
from typing import List, Dict, Any

def generate_document(content_blocks: List[Dict[str, Any]], assumptions: List[str], output_filename: str = "document.docx") -> str:
    """
    Generates a .docx document from the provided content blocks.
    Ensures the output directory exists.
    """
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, output_filename)
    
    doc = Document()
    doc.add_heading("AI Generated Document", 0)
    
    if assumptions:
        doc.add_heading("Assumptions", level=1)
        for assumption in assumptions:
            doc.add_paragraph(f"• {assumption}")
            
    for block in content_blocks:
        if block["type"] == "heading":
            doc.add_heading(block["content"], level=block.get("level", 1))
        elif block["type"] == "paragraph":
            doc.add_paragraph(block["content"])
        elif block["type"] == "table":
            # block["content"] is expected to be a list of lists (rows)
            table_data = block["content"]
            if table_data:
                table = doc.add_table(rows=1, cols=len(table_data[0]))
                table.style = 'Table Grid'
                # Add header row
                hdr_cells = table.rows[0].cells
                for i, col_name in enumerate(table_data[0]):
                    hdr_cells[i].text = str(col_name)
                # Add data rows
                for row_data in table_data[1:]:
                    row_cells = table.add_row().cells
                    for i, cell_val in enumerate(row_data):
                        row_cells[i].text = str(cell_val)
                        
    doc.save(file_path)
    return file_path
